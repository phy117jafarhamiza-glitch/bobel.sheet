import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
import copy
import random
from io import BytesIO
import re

st.set_page_config(page_title="منصة الامتحانات", layout="centered")
st.title("نظام توليد الأسئلة الامتحانية")

TEMPLATE_FILE = 'نموذج الاسئلة 30سؤال.docx' 

# --- دوال مساعدة لتكرار الصفوف (التوسيع التلقائي) ---
def add_row_copy(table, row_idx):
    """تقوم بنسخ صف محدد وإضافته لآخر الجدول"""
    row_copy = copy.deepcopy(table.rows[row_idx]._tr)
    table._tbl.append(row_copy)

def expand_mcq_table(table, current_slots, target_slots):
    """توسيع جدول الاختياري (ينسخ سطرين: سؤال + خيارات)"""
    needed = target_slots - current_slots
    if needed > 0:
        # نفترض أن آخر سطرين هما (سؤال + خيارات)
        # ننسخ آخر صفين ونكررهم
        last_q_row_idx = len(table.rows) - 2
        last_opt_row_idx = len(table.rows) - 1
        
        for _ in range(needed):
            add_row_copy(table, last_q_row_idx) # نسخ سطر السؤال
            add_row_copy(table, last_opt_row_idx) # نسخ سطر الخيارات

def expand_tf_table(table, current_slots, target_slots):
    """توسيع جدول الصح والخطأ (ينسخ سطر واحد)"""
    needed = target_slots - current_slots
    if needed > 0:
        last_row_idx = len(table.rows) - 1
        for _ in range(needed):
            add_row_copy(table, last_row_idx)

# --- الدوال الأساسية ---
def set_document_font_size(doc, size):
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(size)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(size)

def read_questions(file):
    doc = Document(file)
    mcq_list = []
    tf_list = []
    current_mode = None
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        if "# اختياري" in line:
            current_mode = "MCQ"
            i += 1; continue
        elif "# صح وخطأ" in line:
            current_mode = "TF"
            i += 1; continue
            
        if current_mode == "MCQ":
            if i + 5 < len(lines):
                q = lines[i]
                opts = lines[i+1:i+6]
                if not any("#" in opt for opt in opts):
                    mcq_list.append({"q": q, "opts": opts})
                    i += 6; continue
        elif current_mode == "TF":
            tf_list.append(line)
            i += 1; continue
        i += 1
    return mcq_list, tf_list

def generate_exam(mcq_data, tf_data, template_path, target_mcq_count, target_tf_count):
    doc = Document(template_path)
    
    # خلط الأسئلة
    random.shuffle(mcq_data)
    random.shuffle(tf_data)
    
    # قص العدد المطلوب فقط
    final_mcq = mcq_data[:target_mcq_count]
    final_tf = tf_data[:target_tf_count]
    
    mcq_idx = 0
    tf_idx = 0
    current_shuffled_opts = None 
    
    # --- المرحلة 1: توسيع الجداول لتكفي العدد المطلوب ---
    for table in doc.tables:
        row_text_sample = ""
        try:
            for row in table.rows[:2]:
                for cell in row.cells: row_text_sample += cell.text
        except: pass

        # جدول الاختياري
        if "A" in row_text_sample and ("B" in row_text_sample or "," in row_text_sample):
            # نحسب عدد الأسئلة الموجودة حالياً في الجدول
            # كل سؤال يأخذ سطرين (سؤال + خيارات)
            # سنقوم بعد الصفوف التي تحتوي "A" (صفوف الخيارات)
            current_slots = 0
            for row in table.rows:
                if "A" in "".join([c.text for c in row.cells]):
                    current_slots += 1
            
            # التوسيع
            if target_mcq_count > current_slots:
                expand_mcq_table(table, current_slots, target_mcq_count)

        # جدول الصح والخطأ
        else:
            is_tf = False
            for row in table.rows:
                if "(" in "".join([c.text for c in row.cells]): is_tf = True; break
            
            if is_tf:
                current_slots = 0
                for row in table.rows:
                    if "(" in "".join([c.text for c in row.cells]):
                        current_slots += 1
                
                # التوسيع
                if target_tf_count > current_slots:
                    expand_tf_table(table, current_slots, target_tf_count)

    # --- المرحلة 2: التعبئة ---
    for table in doc.tables:
        row_text_sample = ""
        try:
            for row in table.rows[:2]:
                for cell in row.cells: row_text_sample += cell.text
        except: pass

        # الاختياري
        if "A" in row_text_sample and ("B" in row_text_sample or "," in row_text_sample):
            for row in table.rows:
                cells = row.cells
                row_text = "".join([c.text for c in cells])
                
                if "..." in row_text and "A" not in row_text:
                    if mcq_idx < len(final_mcq):
                        current_opts = final_mcq[mcq_idx]['opts']
                        current_shuffled_opts = list(current_opts)
                        random.shuffle(current_shuffled_opts)
                        q_text = final_mcq[mcq_idx]['q']
                        for cell in cells:
                            for p in cell.paragraphs:
                                if "..." in p.text:
                                    p.text = re.sub(r'\.{3,}', q_text, p.text)
                
                elif "A" in row_text and current_shuffled_opts:
                    opt_map = {'A': current_shuffled_opts[0], 'B': current_shuffled_opts[1], 'C': current_shuffled_opts[2], 'D': current_shuffled_opts[3], 'E': current_shuffled_opts[4]}
                    for i in range(len(cells)):
                        cell_text = cells[i].text.strip().replace(",", "")
                        if cell_text in opt_map:
                            if i + 1 < len(cells):
                                next_cell = cells[i+1]
                                next_cell.text = opt_map[cell_text]
                                for p in next_cell.paragraphs: p.alignment = 2 
                    mcq_idx += 1
                    current_shuffled_opts = None

        # الصح والخطأ
        else:
            is_tf = False
            for row in table.rows:
                if "(" in "".join([c.text for c in row.cells]): is_tf = True; break
            if is_tf:
                for row in table.rows:
                    if tf_idx < len(final_tf):
                        full_row = "".join([c.text for c in row.cells])
                        if "..." in full_row and "(" in full_row:
                             for cell in row.cells:
                                for p in cell.paragraphs:
                                    if "..." in p.text:
                                        p.text = re.sub(r'\.{3,}', final_tf[tf_idx], p.text)
                             tf_idx += 1

    set_document_font_size(doc, 10)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- الواجهة ---
st.sidebar.header("لوحة التحكم")
uploaded_file = st.file_uploader("1. ارفع ملف بنك الأسئلة", type=['docx'])

if uploaded_file:
    all_mcq, all_tf = read_questions(uploaded_file)
    if not all_mcq and not all_tf:
        st.error("لم يتم العثور على أسئلة!")
    else:
        st.success(f"المتوفر: {len(all_mcq)} اختياري، {len(all_tf)} صح وخطأ.")
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        with col1:
            mcq_count = st.number_input("عدد الاختيارات المطلوب", 0, len(all_mcq), min(20, len(all_mcq)))
        with col2:
            tf_count = st.number_input("عدد الصح والخطأ المطلوب", 0, len(all_tf), min(10, len(all_tf)))
            
        if st.button("توليد الامتحان"):
            try:
                # نرسل الأرقام المطلوبة للدالة لتقوم بالتوسيع
                final_file = generate_exam(all_mcq, all_tf, TEMPLATE_FILE, mcq_count, tf_count)
                st.download_button("📥 تحميل الامتحان", final_file, "Exam_Expanded.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.success("تم توسيع القالب وتعبئة الأسئلة بنجاح!")
            except Exception as e:
                st.error(f"خطأ: {e}")
